"""Render resume templates with provided data."""
from typing import Dict
import os
import re
import tempfile

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    Document = None
    HAS_DOCX = False

try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
except ImportError:
    DocxTemplate = None
    HAS_DOCXTPL = False


def render_template(template_path: str, data: Dict[str, object], output_path: str) -> str:
    """Render template with data and write to output file.
    
    Supports DOCX templates with advanced formatting preservation via docxtpl,
    or falls back to simple text replacement for other formats.
    
    Args:
        template_path: Path to template file
        data: Dictionary of placeholder keys and values
        output_path: Path where rendered output will be saved
        
    Returns:
        Path to the generated file
        
    Raises:
        RuntimeError: If required dependencies are missing
        FileNotFoundError: If template file doesn't exist
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")
    
    ext = os.path.splitext(template_path)[1].lower()
    
    if ext == ".docx":
        return _render_docx_template(template_path, data, output_path)
    else:
        return _render_text_template(template_path, data, output_path)


def _render_docx_template(template_path: str, data: Dict[str, object], output_path: str) -> str:
    """Render DOCX template using docxtpl or python-docx."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx is required to render .docx templates")
    
    # Try docxtpl first (best formatting preservation)
    if HAS_DOCXTPL:
        try:
            return _render_with_docxtpl(template_path, data, output_path)
        except Exception as e:
            print(f"Warning: docxtpl failed ({e}), falling back to python-docx")
    
    # Fall back to improved python-docx approach
    return _render_with_docx_fallback(template_path, data, output_path)


def _render_with_docxtpl(template_path: str, data: Dict[str, object], output_path: str) -> str:
    """Render using docxtpl with Jinja2 syntax conversion."""
    if not DocxTemplate:
        raise RuntimeError("docxtpl is required for advanced template rendering")
        
    # Convert {KEY} placeholders to {{ KEY }} for Jinja2
    tmpfd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmpfd)
    
    try:
        _convert_placeholders_to_jinja(template_path, tmp_path)
        
        # Render with docxtpl
        tpl = DocxTemplate(tmp_path)
        
        # Filter out skill data for docxtpl (since we kept skill placeholders as {SKILL_X})
        filtered_data = {k: v for k, v in data.items() if not k.startswith("SKILL_")}
        tpl.render(filtered_data)
        tpl.save(output_path)
        
        # Now handle skill placeholders with custom formatting
        _apply_skill_formatting_to_docx(output_path, data)
        
        return output_path
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def _apply_skill_formatting_to_docx(docx_path: str, data: Dict[str, object]):
    """Apply custom skill formatting to a DOCX file after docxtpl processing."""
    if not Document:
        return
        
    doc = Document(docx_path)
    
    # Process paragraphs
    for p in doc.paragraphs:
        if "{" in p.text:
            _simple_replace_paragraph_text(p, data)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{" in cell.text:
                    for paragraph in cell.paragraphs:
                        if "{" in paragraph.text:
                            _simple_replace_paragraph_text(paragraph, data)
    
    doc.save(docx_path)


def _convert_placeholders_to_jinja(template_path: str, output_path: str):
    """Convert {KEY} placeholders to {{ KEY }} format for Jinja2."""
    if not Document:
        raise RuntimeError("python-docx is required")
        
    doc = Document(template_path)
    placeholder_pattern = re.compile(r"\{([A-Za-z0-9_]+)\}")
    
    def convert_text(text: str) -> str:
        return placeholder_pattern.sub(lambda m: "{{ " + m.group(1) + " }}", text)
    
    # Convert paragraphs (but skip skill placeholders)
    for p in doc.paragraphs:
        if "{" in p.text:
            # Don't convert skill placeholders to Jinja2 format - keep them as {SKILL_X}
            placeholders = re.findall(r"\{([A-Za-z0-9_]+)\}", p.text)
            conversion_dict = {}
            for k in placeholders:
                if not k.startswith("SKILL_"):
                    conversion_dict[k] = "{{ " + k + " }}"
            if conversion_dict:
                _simple_replace_paragraph_text(p, conversion_dict)
    
    # Convert tables (but skip skill placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{" in cell.text:
                    for paragraph in cell.paragraphs:
                        if "{" in paragraph.text:
                            # Don't convert skill placeholders to Jinja2 format
                            placeholders = re.findall(r"\{([A-Za-z0-9_]+)\}", paragraph.text)
                            conversion_dict = {}
                            for k in placeholders:
                                if not k.startswith("SKILL_"):
                                    conversion_dict[k] = "{{ " + k + " }}"
                            if conversion_dict:
                                _simple_replace_paragraph_text(paragraph, conversion_dict)
    
    doc.save(output_path)


def _render_with_docx_fallback(template_path: str, data: Dict[str, object], output_path: str) -> str:
    """Fallback DOCX rendering using simple text replacement."""
    if not Document:
        raise RuntimeError("python-docx is required")
        
    doc = Document(template_path)
    
    # Simple approach: replace in all text, then worry about formatting later
    for p in doc.paragraphs:
        if "{" in p.text:
            _simple_replace_paragraph_text(p, data)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{" in cell.text:
                    for paragraph in cell.paragraphs:
                        if "{" in paragraph.text:
                            _simple_replace_paragraph_text(paragraph, data)
    
    doc.save(output_path)
    return output_path


def _simple_replace_paragraph_text(paragraph, data):
    """Replace text while preserving formatting for job titles, locations, and skills."""
    # Get the full text from all runs
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Check if this looks like a job title line with location
    has_job_title = any(f"{{{key}}}" in full_text for key in data.keys() if "TITLE" in key.upper())
    has_location_pattern = " - " in full_text or " – " in full_text
    
    # Check if this looks like a skill placeholder
    has_skill_placeholder = any(f"{{{key}}}" in full_text for key in data.keys() if key.startswith("SKILL_"))
    
    print(f"DEBUG: Processing paragraph: '{full_text}'")
    print(f"DEBUG: has_job_title={has_job_title}, has_location_pattern={has_location_pattern}, has_skill_placeholder={has_skill_placeholder}")
    
    if has_job_title and has_location_pattern:
        print("DEBUG: Using job title formatting")
        _replace_job_title_with_formatting(paragraph, data)
    elif has_skill_placeholder:
        print("DEBUG: Using skill formatting")
        _replace_skill_with_formatting(paragraph, data)
    else:
        print("DEBUG: Using simple text replacement")
        # Use simple replacement for other cases
        _simple_text_replacement(paragraph, data, full_text)


def _replace_job_title_with_formatting(paragraph, data):
    """Replace job title while preserving bold title and italic location."""
    from docx.shared import Pt
    
    # Get the full text
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Replace placeholders
    new_text = full_text
    job_title_replacement = ""
    
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        if placeholder in new_text:
            if "TITLE" in key.upper():
                job_title_replacement = str(value)
            new_text = new_text.replace(placeholder, str(value))
    
    # Clear existing runs
    for run in paragraph.runs:
        run.clear()
    
    # Split on the separator to get title and location parts
    if " – " in new_text:
        separator = " – "
    elif " - " in new_text:
        separator = " - "
    else:
        # Fallback to simple replacement
        if paragraph.runs:
            run = paragraph.runs[0]
            run.text = new_text
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        return
    
    parts = new_text.split(separator, 1)
    if len(parts) == 2:
        title_part = parts[0].strip()
        location_part = parts[1].strip()
        
        # Add title part (bold, Times New Roman 11)
        title_run = paragraph.add_run(title_part)
        title_run.bold = True
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        
        # Add separator (Times New Roman 11)
        sep_run = paragraph.add_run(separator)
        sep_run.font.name = 'Times New Roman'
        sep_run.font.size = Pt(11)
        
        # Add location part (italic, Times New Roman 11)
        location_run = paragraph.add_run(location_part)
        location_run.italic = True
        location_run.font.name = 'Times New Roman'
        location_run.font.size = Pt(11)
    else:
        # Fallback
        if paragraph.runs:
            run = paragraph.runs[0]
            run.text = new_text
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)


def _replace_skill_with_formatting(paragraph, data):
    """Replace skill placeholder with formatted text: Category[Skill, Skill, Skill]."""
    from docx.shared import Pt
    
    # Get the full text
    full_text = "".join(run.text for run in paragraph.runs)
    print(f"DEBUG: _replace_skill_with_formatting called with full_text: '{full_text}'")
    
    # Find and replace skill placeholders (handle both {KEY} and {{ KEY }} formats)
    new_text = full_text
    for key, value in data.items():
        if key.startswith("SKILL_"):
            # Try both placeholder formats
            placeholder1 = f"{{{key}}}"  # {SKILL_1}
            placeholder2 = f"{{{{ {key} }}}}"  # {{ SKILL_1 }}
            
            if placeholder1 in new_text:
                print(f"DEBUG: Replacing {placeholder1} with '{value}'")
                new_text = new_text.replace(placeholder1, str(value))
            elif placeholder2 in new_text:
                print(f"DEBUG: Replacing {placeholder2} with '{value}'")
                new_text = new_text.replace(placeholder2, str(value))
    
    # Clear existing runs
    for run in paragraph.runs:
        run.clear()
    
    # Parse the skill format: Category [Skill, Skill, Skill]
    skill_pattern = re.compile(r'^([^[]+?)\s*\[([^\]]+)\]$')
    match = skill_pattern.match(new_text.strip())
    
    if match:
        category = match.group(1).strip()
        skills = match.group(2).strip()
        
        # Add category part (bold, Times New Roman 11)
        category_run = paragraph.add_run(category)
        category_run.bold = True
        category_run.font.name = 'Times New Roman'
        category_run.font.size = Pt(11)
        
        # Add space
        space_run = paragraph.add_run(' ')
        space_run.font.name = 'Times New Roman'
        space_run.font.size = Pt(11)
        
        # Add opening bracket
        bracket_run = paragraph.add_run('[')
        bracket_run.font.name = 'Times New Roman'
        bracket_run.font.size = Pt(11)
        
        # Add skills part (normal, Times New Roman 11)
        skills_run = paragraph.add_run(skills)
        skills_run.font.name = 'Times New Roman'
        skills_run.font.size = Pt(11)
        
        # Add closing bracket
        bracket_run2 = paragraph.add_run(']')
        bracket_run2.font.name = 'Times New Roman'
        bracket_run2.font.size = Pt(11)
    else:
        # Fallback to simple text
        run = paragraph.add_run(new_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)


def _simple_text_replacement(paragraph, data, full_text):
    """Simple text replacement without formatting preservation."""
    new_text = full_text
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        new_text = new_text.replace(placeholder, str(value))
    
    # If text changed, update the paragraph
    if new_text != full_text:
        # Clear all runs and create a single run with the new text
        for run in paragraph.runs:
            run.clear()
        
        # Add the new text to the first run
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def _render_text_template(template_path: str, data: Dict[str, object], output_path: str) -> str:
    """Render text-based template with simple replacement."""
    with open(template_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    for k, v in data.items():
        content = content.replace(f"{{{{{k}}}}}", str(v))
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    return output_path